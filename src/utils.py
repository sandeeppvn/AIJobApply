import logging
import os
import re

import docx
import docxtpl
import pandas as pd
from dotenv import find_dotenv, load_dotenv

from src.google_drive_handler import GoogleDriveHandler

# Setting up logger
logger = logging.getLogger(__name__)

TEMPLATES = ["cover_letter_template", "resume_template", "email_template", "linkedin_note_template"]

def get_file_content(path: str) -> str:
    """Get the content of a file."""
    try:
        # Ensure the file is .docx only
        if not path.endswith((".docx")):
            raise ValueError("Only .docx files are supported.")

        # Read the entire docx file as a string
        doc = docx.Document(path)
        content = "".join([paragraph.text for paragraph in doc.paragraphs])
        return content
    except Exception as e:
        logger.exception(f"Error reading file: {e}")
        return ""


def create_job_folder(job:pd.Series, resume_path:str, google_drive_handler:GoogleDriveHandler, destination:str = "Job Applications"):
    """
    Create a folder for the job application process.
    Add relevant files to the folder.
    """
    # Create a folder for the job application
    # Name the folder as CompanyName_Position, no special characters or spaces. Add an underscore between company name and position
    company_name = re.sub(r"[^\w\s]", "", job["Company Name"])
    position = re.sub(r"\W+", "", job["Position"])
    job_name = f"{company_name}_{position}"
    job_folder = f"{destination}/{job_name}"
    if os.path.exists(job_folder):
        import shutil
        shutil.rmtree(job_folder)
    os.makedirs(job_folder, exist_ok=True)

    # Get the ID of the job folder (Creating the folder if it doesn't exist)
    job_folder_id = google_drive_handler.get_folder(job_name, google_drive_handler.job_root_folder_id)

    # Load the resume and replace jinja2 variables
    resume_variables = {
        "resume_professional_summary": job["Resume"],
    }
    resume = docxtpl.DocxTemplate(resume_path)
    resume.render(resume_variables)
    resume_file_path = os.path.join(job_folder, "Resume.docx")
    resume.save(resume_file_path)
    google_drive_handler.upload_file("Resume.docx", resume_file_path, job_folder_id)

    cover_letter_text = job["Cover Letter"]
    cover_letter_doc = docx.Document()
    cover_letter_doc.add_paragraph(cover_letter_text)
    cover_letter_file_path = os.path.join(job_folder, "Cover Letter.docx")
    cover_letter_doc.save(cover_letter_file_path)
    google_drive_handler.upload_file("Cover Letter.docx", cover_letter_file_path, job_folder_id)

    email = job["Message Subject"] + "\n\n" + job["Message Content"]
    email_file_path = os.path.join(job_folder, "Email.txt")
    with open(email_file_path, "w", encoding="utf-8") as file:
        file.write(email)
    google_drive_handler.upload_file("Email.txt", email_file_path, job_folder_id)

    linkedin_note = job["LinkedIn Note"]
    linkedin_note_file_path = os.path.join(job_folder, "LinkedIn Note.txt")
    with open(linkedin_note_file_path, "w", encoding="utf-8") as file:
        file.write(linkedin_note)
    google_drive_handler.upload_file("LinkedIn Note.txt", linkedin_note_file_path, job_folder_id)


def validate_arguments(args: dict) -> dict:
    """
    Validate the arguments passed to the CLI. 
    Check if all arguments are present and valid either from the CLI or as an environment variable.
    Create a dictionary of the arguments and their updated values.
    """
    load_dotenv(find_dotenv())
    validate_args = {}

    required_args = {
        # LLM arguments
        "LLM_API_KEY": "LLM api key",
        "LLM_MODEL": "LLM model to use",
        # "LLM_API_URL": "LLM API URL",

        # Google Sheets arguments
        "GOOGLE_API_CREDENTIALS_FILE": "Path to the credentials file for google api",
        "GOOGLE_SHEET_NAME": "Name of the google sheet to read jobs from",

        # Document arguments
        "RESUME_PATH": "Path to resume",
        "RESUME_PROFESSIONAL_SUMMARY": "Professional summary for resume",
        "COVER_LETTER_PATH": "Path to cover letter",
        "DESTINATION_FOLDER": "Folder to save documents to",
        
    }

    gmail_args = {
        "GMAIL_ADDRESS": "Gmail address to send emails from",
        "GMAIL_PASSWORD": "Password to gmail account",
    }

    linkedin_args = {
        "CHROMEDRIVER_PATH": "Path to the selenium driver",
        "LINKEDIN_USERNAME": "LinkedIn username",
        "LINKEDIN_PASSWORD": "LinkedIn password",
        # "LINKEDIN_NOTE": "LinkedIn note",
    }

    # Check if all required arguments are provided
    for arg_name, arg_description in required_args.items():
        if arg_name not in args and os.getenv(arg_name) is None:
            raise ValueError(f"Required argument '{arg_name}' not provided.")
        validate_args[arg_name] = args[arg_name] if arg_name in args else os.getenv(arg_name)
        
    # Check if all Gmail arguments are provided
    if "USE_GMAIL" in args and args["USE_GMAIL"]:
        validate_args["USE_GMAIL"] = True
        for arg_name, arg_description in gmail_args.items():
            if arg_name not in args and os.getenv(arg_name) is None:
                raise ValueError(f"Gmail argument '{arg_name}' not provided.")
            # Validate the email is a valid email address
            if arg_name == "GMAIL_ADDRESS":
                from validate_email import validate_email
                if not validate_email(args[arg_name] if arg_name in args else os.getenv(arg_name)):
                    raise ValueError(f"Invalid email address: {args[arg_name] if arg_name in args else os.getenv(arg_name)}")
            validate_args[arg_name] = args[arg_name] if arg_name in args else os.getenv(arg_name)
    else:
        validate_args["USE_GMAIL"] = False
            
    # Check if all LinkedIn arguments are provided
    if "USE_LINKEDIN" in args and args["USE_LINKEDIN"]:
        validate_args["USE_LINKEDIN"] = True
        for arg_name, arg_description in linkedin_args.items():
            if arg_name not in args and os.getenv(arg_name) is None:
                raise ValueError(f"LinkedIn argument '{arg_name}' not provided.")
            validate_args[arg_name] = args[arg_name] if arg_name in args else os.getenv(arg_name)

        if args["INTERACTIVE"]:
            validate_args["INTERACTIVE"] = True
        else:
            validate_args["INTERACTIVE"] = False
    else:
        validate_args["USE_LINKEDIN"] = False
        validate_args["INTERACTIVE"] = False

    # validate_args["LLM_API_URL"] = validate_args["LLM_API_URL"].strip()
    # if not validate_args["LLM_API_URL"].startswith("https://"):
    #     raise ValueError(f"LLM_API_URL must start with 'https://'.")
        
    

    return validate_args